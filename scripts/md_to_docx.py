"""
Markdown → Word (.docx) 转换器

将采访转录技能生成的 Markdown 文档转换为排版清晰的 Word 文档。
专为 interview-transcriber 的输出格式定制，支持：
  - 标题 (# / ##)
  - 居中人物静帧（<div align="center"><img .../></div>）
  - 引用块元数据（> 行）
  - 人物信息表格（| 字段 | 内容 |）
  - 说话人标签行（**采访者** [MM:SS]）
  - 行内加粗（**text**）
  - 无序/有序列表（- / 1.）

用法:
  python md_to_docx.py <input.md> [output.docx]
  不指定 output.docx 时，自动将扩展名替换为 .docx（与原 md 同目录）
"""

import os
import sys
import re
import argparse

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误: python-docx 未安装。请先执行: pip install python-docx")
    sys.exit(1)


def add_inline_runs(paragraph, text):
    """将含 **加粗** 的文本拆分为带格式 run 添加到段落。"""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:  # 奇数段为加粗
            run.bold = True


def parse_px(width_str):
    """解析 '280' 或 '280px' 为 Inches。"""
    try:
        px = float(re.sub(r"[^0-9.]", "", width_str or "280"))
    except ValueError:
        px = 280.0
    return Inches(px / 96.0)


def is_table_row(line):
    line = line.strip()
    return line.startswith("|") and line.endswith("|") and line.count("|") >= 2


def is_separator_row(line):
    s = line.strip().strip("|").replace(" ", "")
    return bool(re.fullmatch(r":?-+:?", s)) and "-" in s


def split_table_row(line):
    # 去掉首尾的 "|"，按 "|" 切分，并去空白
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def convert(md_path, docx_path):
    base_dir = os.path.dirname(os.path.abspath(md_path))

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()

    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        line = raw.rstrip("\n")
        stripped = line.strip()

        # 跳过空行
        if stripped == "":
            i += 1
            continue

        # 跳过 HTML 包裹标签（静帧块的一部分）
        if stripped.startswith("<div") or stripped == "</div>":
            i += 1
            continue

        # 分隔线 ---
        if stripped == "---":
            i += 1
            continue

        # 居中静帧：<div align="center"> ... <img .../> ... </div>
        if "<img" in line:
            src_match = re.search(r'src="([^"]+)"', line)
            width_match = re.search(r'width="([^"]+)"', line)
            if src_match:
                img_src = src_match.group(1)
                img_path = img_src if os.path.isabs(img_src) else os.path.join(base_dir, img_src)
                width = parse_px(width_match.group(1)) if width_match else Inches(280 / 96.0)
                if os.path.exists(img_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        p.add_run().add_picture(img_path, width=width)
                    except Exception:
                        # 直接插入失败（如扩展名与真实格式不符），用 Pillow 重新编码为 JPEG 再插入
                        try:
                            from io import BytesIO
                            from PIL import Image
                            buf = BytesIO()
                            with Image.open(img_path) as im:
                                im.convert("RGB").save(buf, "JPEG")
                            buf.seek(0)
                            p.add_run().add_picture(buf, width=width)
                        except Exception as e2:
                            print(f"  ⚠️ 插入图片失败: {e2}")
                else:
                    print(f"  ⚠️ 静帧图片不存在，跳过: {img_path}")
            i += 1
            continue

        # 标题
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        # 表格
        if is_table_row(line):
            # 收集连续表格行
            rows = []
            while i < n and is_table_row(lines[i]):
                rows.append(lines[i].strip())
                i += 1
            # 过滤分隔行
            data_rows = [r for r in rows if not is_separator_row(r)]
            if data_rows:
                ncols = max(len(split_table_row(r)) for r in data_rows)
                table = doc.add_table(rows=1, cols=ncols)
                try:
                    table.style = "Table Grid"
                except Exception:
                    pass
                # 首行
                for idx, cell_text in enumerate(split_table_row(data_rows[0])):
                    add_inline_runs(table.rows[0].cells[idx].paragraphs[0], cell_text)
                # 其余行
                for r in data_rows[1:]:
                    cells = split_table_row(r)
                    row = table.add_row().cells
                    for idx, cell_text in enumerate(cells):
                        if idx < len(row):
                            add_inline_runs(row[idx].paragraphs[0], cell_text)
            continue

        # 引用块（元数据 / 占位说明）
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                q = lines[i].strip()
                q = re.sub(r"^>\s?", "", q)
                quote_lines.append(q)
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            # 逐行，行间换行
            for j, ql in enumerate(quote_lines):
                if j > 0:
                    p.add_run().add_break()
                add_inline_runs(p, ql)
            continue

        # 列表
        if re.match(r"^-\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, re.sub(r"^-\s+", "", stripped))
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # 说话人标签行：**采访者** [MM:SS]
        speaker_match = re.match(r"^\*\*(.+?)\*\*\s*\[.*\]$", stripped)
        if speaker_match:
            p = doc.add_paragraph()
            add_inline_runs(p, stripped)
            i += 1
            continue

        # 普通段落（合并后续连续非空、非特殊行）
        buf = [stripped]
        i += 1
        while i < n and lines[i].strip() != "" and not (
            lines[i].strip().startswith("#")
            or lines[i].strip().startswith(">")
            or lines[i].strip().startswith("---")
            or is_table_row(lines[i])
            or "<img" in lines[i]
        ):
            buf.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(buf))

    doc.save(docx_path)
    print(f"✅ Word 文档已生成: {docx_path}")


def main():
    parser = argparse.ArgumentParser(description="Markdown → Word (.docx) 转换器")
    parser.add_argument("input", help="输入的 Markdown 文件路径")
    parser.add_argument("output", nargs="?", default=None, help="输出 .docx 路径（默认同目录替换扩展名）")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误: 输入文件不存在: {args.input}")
        sys.exit(1)

    out = args.output or os.path.splitext(args.input)[0] + ".docx"
    convert(args.input, out)


if __name__ == "__main__":
    main()
