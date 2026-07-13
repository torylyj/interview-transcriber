"""
结构化文档 JSON → Word (.docx) 生成器

直接读取 interview-transcriber 流程产出的 <标题>_document.json（Agent 完成
说话人识别 / 摘要生成 / 语气词自检后写入的最终结构化数据），生成排版清晰的
Word 文档。**全程不依赖 Markdown 中间文件。**

支持：
  - 文档标题
  - 居中人物静帧（仅视频输入，frame_path 非 null 时）
  - 内容摘要（📝 内容摘要，多段落）
  - 受访人信息表格（👤 受访人信息，2 列）
  - 文档信息引用块（📋 文档信息）
  - 采访记录（💬 采访记录，加粗说话人标签 + 时间码）

可选：
  --export-md <path>  额外导出一份临时 Markdown（供在线平台上传用，上传后即删）

用法:
  python build_docx.py <document.json> <output.docx> [--export-md _upload.md]
"""

import os
import sys
import json
import argparse
import re


def normalize_path(p):
    """将 Git Bash 风格路径 /c/Users/... 归一化为 Windows 路径 C:/Users/...。"""
    if not p:
        return p
    m = re.match(r"^/([a-zA-Z])/(.*)$", p)
    if m:
        return f"{m.group(1).upper()}:/{m.group(2)}"
    return p


try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("错误: python-docx 未安装。请先执行: pip install python-docx")
    sys.exit(1)


def add_inline_runs(paragraph, text):
    """将含 **加粗** 的文本拆分为带格式 run 添加到段落。"""
    if text is None:
        return
    parts = str(text).split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:  # 奇数段为加粗
            run.bold = True


def parse_px(width_str):
    """解析 '280' 或 '280px' 为 Inches。"""
    try:
        px = float("".join(ch for ch in (width_str or "280") if ch.isdigit() or ch == "."))
    except ValueError:
        px = 280.0
    if px <= 0:
        px = 280.0
    return Inches(px / 96.0)


def add_image(doc, img_path, base_dir, width=Inches(280 / 96.0)):
    """居中插入图片，失败时用 Pillow 重新编码后重试。"""
    img_path = img_path if os.path.isabs(img_path) else os.path.join(base_dir, img_path)
    if not os.path.exists(img_path):
        print(f"  ⚠️ 静帧图片不存在，跳过: {img_path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(img_path, width=width)
        return
    except Exception:
        pass
    # 直接插入失败（扩展名与真实格式不符等），用 Pillow 重新编码为 JPEG 再插入
    try:
        from io import BytesIO
        from PIL import Image
        buf = BytesIO()
        with Image.open(img_path) as im:
            im.convert("RGB").save(buf, "JPEG")
        buf.seek(0)
        p.add_run().add_picture(buf, width=width)
    except Exception as e2:
        print(f"  ⚠️ 插入图片失败: {e2}")


def build(doc, data, base_dir):
    # 标题
    doc.add_heading(data.get("title", "转录文档"), level=0)

    # 居中静帧（仅视频输入）
    frame_path = data.get("frame_path")
    if frame_path:
        add_image(doc, frame_path, base_dir, parse_px("280"))

    doc.add_paragraph("")

    # ── 内容摘要 ──
    summary = data.get("summary")
    if summary:
        doc.add_heading("\U0001F4DD 内容摘要", level=1)
        for para in str(summary).split("\n"):
            if para.strip() == "":
                doc.add_paragraph("")
            else:
                p = doc.add_paragraph()
                add_inline_runs(p, para)
        doc.add_paragraph("")

    # ── 受访人信息（条件性：无信息整段省略；多人每人一个表格）──
    person_info = data.get("person_info") or []
    if person_info:
        # 兼容两种结构：
        #   新结构 = [{"name": "...", "fields": [{field,value}]}, ...]（支持多人）
        #   旧结构 = [{"field": "...", "value": "..."}, ...]（单人，整体当一个人的字段）
        new_style = isinstance(person_info[0], dict) and ("fields" in person_info[0])
        people = person_info if new_style else [{"name": "", "fields": person_info}]
        doc.add_heading("\U0001F464 受访人信息", level=1)
        for person in people:
            name = person.get("name", "") if isinstance(person, dict) else ""
            fields = person.get("fields", []) if isinstance(person, dict) else []
            if not fields:
                continue
            if name:
                pname = doc.add_paragraph()
                pname.add_run(name).bold = True
            table = doc.add_table(rows=1, cols=2)
            try:
                table.style = "Table Grid"
            except Exception:
                pass
            # 表头
            hdr = table.rows[0].cells
            hdr[0].paragraphs[0].add_run("字段").bold = True
            hdr[1].paragraphs[0].add_run("内容").bold = True
            for item in fields:
                row = table.add_row().cells
                add_inline_runs(row[0].paragraphs[0], item.get("field", ""))
                add_inline_runs(row[1].paragraphs[0], item.get("value", ""))
            doc.add_paragraph("")

    # ── 文档信息 ──
    doc.add_heading("\U0001F4CB 文档信息", level=1)
    tool = data.get("transcription_tool", "")
    tc_accuracy = "段内为估算值（4分钟粒度）" if ("Qwen" in tool or "云端" in tool) else "句级时间码为文本长度插值估算（段落边界精确）"
    info_lines = [
        f"源文件：{data.get('source_file', '')}",
        f"输入类型：{data.get('input_type', '')}",
        f"转录工具：{tool}",
        f"说话人识别：{data.get('speaker_method', 'LLM 语义分析')}",
        f"摘要与人物信息：{data.get('summary_method', 'LLM 生成')}",
        f"转录日期：{data.get('date', '')}",
        f"时间码精度：{tc_accuracy}",
    ]
    q = doc.add_paragraph()
    q.paragraph_format.left_indent = Inches(0.3)
    for j, line in enumerate(info_lines):
        if j > 0:
            q.add_run().add_break()
        q.add_run(line)
    doc.add_paragraph("")

    # ── 采访记录 ──
    doc.add_heading("\U0001F4AC 采访记录", level=1)
    conversation = data.get("conversation") or []
    if not conversation:
        print("  ⚠️ 警告: conversation 为空，文档将缺少采访记录")
    for turn in conversation:
        speaker = turn.get("speaker", "")
        timestamp = turn.get("timestamp", "")
        text = turn.get("text", "")
        # 说话人 + 时间码行（加粗说话人）
        label = f"**{speaker}** {timestamp}".strip()
        p_label = doc.add_paragraph()
        add_inline_runs(p_label, label)
        # 正文
        p_text = doc.add_paragraph()
        add_inline_runs(p_text, text)


def export_markdown(data, md_path):
    """将结构化文档导出为临时 Markdown（供在线平台上传；上传后即删）。"""
    lines = [f"# {data.get('title', '转录文档')}", ""]

    frame_path = data.get("frame_path")
    if frame_path:
        lines += [f"![人物静帧]({frame_path})", ""]

    lines += ["---", "", "## \U0001F4DD 内容摘要", "", data.get("summary", ""), "", "---", ""]
    person_info = data.get("person_info") or []
    if person_info:
        lines.append("## \U0001F464 受访人信息")
        new_style = isinstance(person_info[0], dict) and ("fields" in person_info[0])
        people = person_info if new_style else [{"name": "", "fields": person_info}]
        for person in people:
            name = person.get("name", "") if isinstance(person, dict) else ""
            fields = person.get("fields", []) if isinstance(person, dict) else []
            if not fields:
                continue
            if name:
                lines.append(f"### {name}")
            lines += ["", "| 字段 | 内容 |", "|------|------|"]
            for item in fields:
                lines.append(f"| {item.get('field', '')} | {item.get('value', '')} |")
            lines.append("")
    lines += ["", "---", "", "## \U0001F4CB 文档信息", ""]
    tool = data.get("transcription_tool", "")
    tc_accuracy = "段内为估算值（4分钟粒度）" if ("Qwen" in tool or "云端" in tool) else "句级时间码为文本长度插值估算（段落边界精确）"
    lines += [
        f"> 源文件：{data.get('source_file', '')}",
        f"> 输入类型：{data.get('input_type', '')}",
        f"> 转录工具：{tool}",
        f"> 说话人识别：{data.get('speaker_method', 'LLM 语义分析')}",
        f"> 摘要与人物信息：{data.get('summary_method', 'LLM 生成')}",
        f"> 转录日期：{data.get('date', '')}",
        f"> 时间码精度：{tc_accuracy}",
        "",
        "---",
        "",
        "## \U0001F4AC 采访记录",
        "",
    ]
    for turn in data.get("conversation") or []:
        speaker = turn.get("speaker", "")
        timestamp = turn.get("timestamp", "")
        text = turn.get("text", "")
        lines.append(f"**{speaker}** {timestamp}")
        lines.append(text)
        lines.append("")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"✅ 临时 Markdown 已导出（上传后请删除）: {md_path}")


def main():
    parser = argparse.ArgumentParser(description="结构化文档 JSON → Word (.docx) 生成器")
    parser.add_argument("input", help="输入的文档 JSON 文件路径（_document.json）")
    parser.add_argument("output", help="输出 .docx 路径")
    parser.add_argument("--export-md", default=None, help="额外导出临时 Markdown 的路径（可选）")
    args = parser.parse_args()

    # 归一化路径（兼容 Git Bash 的 /c/... 写法，否则 Windows 原生程序找不到文件）
    args.input = normalize_path(args.input)
    args.output = normalize_path(args.output)
    if args.export_md:
        args.export_md = normalize_path(args.export_md)

    if not os.path.exists(args.input):
        print(f"错误: 输入文件不存在: {args.input}")
        sys.exit(1)

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    base_dir = os.path.dirname(os.path.abspath(args.input))
    doc = Document()
    build(doc, data, base_dir)
    doc.save(args.output)
    print(f"✅ Word 文档已生成: {args.output}")

    if args.export_md:
        export_markdown(data, args.export_md)


if __name__ == "__main__":
    main()
